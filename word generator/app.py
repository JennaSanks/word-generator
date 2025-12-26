# Flask imports for web handling
from flask import Flask, render_template, request, send_file

# python-docx imports for Word file creation
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Standard library imports
import uuid   #  Used to generate unique IDs (avoid filename collisions)
import os     # Used for directory and file path handling
import re     # Used for regex (cleaning filenames)

# Utility function: make filename safe for OS
def safe_filename(title):

    """
    Converts the document title into a safe filename.

    Why needed:
    - Windows/macOS/Linux disallow characters like / \ : * ? " < > |
    - Spaces can cause issues in URLs

    What this does:
    - Trims whitespace
    - Removes illegal characters
    - Replaces spaces with underscores
    """

    title = title.strip()
    title = re.sub(r'[\\/:*?"<>|]', '', title)
    title = title.replace(" ", "_")
    return title or "Document" # Fallback name if title is empty

# Fallback name if title is empty
app = Flask(__name__)

# Directory where generated Word files are stored
GENERATED_DIR = "generated"
os.makedirs(GENERATED_DIR, exist_ok=True) # Create if not exists

# Home route – serves the UI
@app.route("/")
def index():
    return render_template("index.html")

# Document generation route
@app.route("/generate", methods=["POST"])
def generate_doc():

    """
    Receives JSON data from frontend,
    creates a Word (.docx) file,
    saves it on the server,
    returns filename for download.
    """

    # Parse JSON sent from frontend
    data = request.json
    title = data["title"]
    sections = data["sections"]

    # Create a new Word document
    doc = Document()

    # Margins
    # Page layout configuration
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    # Title formatting
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add spacing after title
    doc.add_paragraph("")

    # Content rendering
    for section in sections:
        # Section heading
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section["heading"])
        heading_run.bold = True
        heading_run.font.size = Pt(14)

        # Bullet points
        for point in section["points"]:
            bullet = doc.add_paragraph(point, style="List Bullet")
            bullet.paragraph_format.left_indent = Inches(0.25)

        # Space after each section
        doc.add_paragraph("")

    # File naming logic
    safe_title = safe_filename(title)

    # Unique ID avoids overwriting files with same title
    unique_id = uuid.uuid4().hex[:6]
    filename = f"{safe_title}_{unique_id}.docx"

    path = os.path.join(GENERATED_DIR, filename)

    # Save Word document
    doc.save(path)

    # Return filename to frontend for download
    return {"file": filename}


# Secure download route
from werkzeug.utils import secure_filename

@app.route("/download/<filename>")
def download(filename):
    """
    Sends the generated Word file to the user.

    secure_filename():
    - Prevents path traversal attacks
    - Ensures filename is OS-safe
    """
    filename = secure_filename(filename)
    return send_file(
        os.path.join(GENERATED_DIR, filename),
        as_attachment=True
    )

# App entry point
if __name__ == "__main__":
    """
    Runs the Flask development server.
    debug=True enables auto-reload and error tracing.
    """
    app.run(debug=True)


'''
HOW TO RUN LOCALLY:
pip install -r requirements.txt
python app.py


Open:
http://127.0.0.1:5000
'''